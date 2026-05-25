from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
from docx import Document as DocxDocument
import re


def load_document(file_path: str) -> list:
    if file_path.endswith(".pdf"):
        loader = PyPDFLoader(file_path)
    elif file_path.endswith(".txt") or file_path.endswith(".md"):
        loader = TextLoader(file_path)
    elif file_path.endswith(".docx"):
        doc = DocxDocument(file_path)
        parts = []
        parts.extend([p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = (cell.text or "").strip()
                    if t:
                        parts.append(t)
        text = "\n".join(parts)
        return [Document(page_content=text, metadata={"source": file_path})]
    else:
        raise ValueError(f"不支持的文件格式: {file_path}")

    return loader.load()


def _markdown_sections(document: Document) -> list[Document]:
    text = document.page_content or ""
    heading_re = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
    sections = []
    heading_stack: list[str] = []
    current_lines: list[str] = []
    current_title = ""

    def flush():
        if not current_lines:
            return
        content = "\n".join(current_lines).strip()
        if not content:
            return
        title = current_title or "正文"
        page_content = f"章节：{title}\n\n{content}"
        metadata = dict(document.metadata or {})
        metadata["section_title"] = title
        sections.append(Document(page_content=page_content, metadata=metadata))

    for line in text.splitlines():
        m = heading_re.match(line)
        if m:
            flush()
            level = len(m.group(1))
            title = m.group(2).strip()
            heading_stack[:] = heading_stack[: level - 1]
            heading_stack.append(title)
            current_title = " > ".join(heading_stack)
            current_lines = [line]
            continue
        current_lines.append(line)

    flush()
    return sections or [document]


def split_documents(documents: list, chunk_size: int = 700, chunk_overlap: int = 150) -> list:
    source_docs = []
    for doc in documents:
        source = str((doc.metadata or {}).get("source") or "")
        if source.endswith(".md") or re.search(r"^#{1,6}\s+", doc.page_content or "", re.MULTILINE):
            source_docs.extend(_markdown_sections(doc))
        else:
            source_docs.append(doc)

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ".", "。", " ", ""],
    )
    chunks = text_splitter.split_documents(source_docs)
    for chunk in chunks:
        title = (chunk.metadata or {}).get("section_title")
        if title and not chunk.page_content.startswith("章节："):
            chunk.page_content = f"章节：{title}\n\n{chunk.page_content}"
    return chunks
